from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
import os
from tkinter import *
import sqlite3
import tkinter.ttk as ttk

def bilgi_girişi(event):
    liste1=liste.get(ACTIVE)

    okul_adi.delete(0,END)
    okul_ilce_il.delete(0,END)
    egitim_ogretim_yili.delete(0,END)
    bugunun_tarihi.delete(0,END)
    adi_soyadi.delete(0,END)
    adres.delete(0,END)
    adres_ilce.delete(0,END)
    adres_il.delete(0,END)    
    belge_no.delete(0,END)
    tc_no.delete(0,END)
    dogum_yeri.delete(0,END)
    baba_adi.delete(0,END)
    kizi_oglu.delete(0,END)
    diploma_derece.delete(0,END)
    diploma_tarihi.delete(0,END)
    diploma_sayisi.delete(0,END)
    duzenleyen_adi_soyadi.delete(0,END)
    duzenleyen_unvan.delete(0,END)
    onaylayan_adi_soyadi.delete(0,END)
    onaylayan_unvan.delete(0,END)
                   
    vt = sqlite3.connect(str(liste1)+'.sq3')
    im= vt.cursor()
    im.execute(""" SELECT * FROM diploma""")
    rows = im.fetchall()
    data_str = ""
    sf = "{}{}{}{}{}{}{}{}{}{}{}{}{}{}{}{}{}{}{}{}"
    for row in rows:
        data_str += sf.format(row[0], row[1], row[2],row[3], row[4], row[5], row[6], row[7],row[8], row[9], row[10], row[11], row[12],row[13], row[14],row[15], \
                              row[16], row[17],row[18], row[19])

        okul_adi.insert(END,row[0])
        okul_ilce_il.insert(END,row[1])
        egitim_ogretim_yili.insert(END,row[2])
        bugunun_tarihi.insert(END,row[3])
        adi_soyadi.insert(END,row[4])
        adres.insert(END,row[5])
        adres_ilce.insert(END,row[6])
        adres_il.insert(END,row[7]) 
        belge_no.insert(END,row[8])
        tc_no.insert(END,row[9])
        dogum_yeri.insert(END,row[10])
        baba_adi.insert(END,row[11])
        kizi_oglu.insert(END,row[12])
        diploma_derece.insert(END,row[13])
        diploma_tarihi.insert(END,row[14])
        diploma_sayisi.insert(END,row[15])
        duzenleyen_adi_soyadi.insert(END,row[16])
        duzenleyen_unvan.insert(END,row[17])
        onaylayan_adi_soyadi.insert(END,row[18])
        onaylayan_unvan.insert(END,row[19])
                
def kaydet():
    okul_adi1 = okul_adi.get()
    okul_ilce_il1 = okul_ilce_il.get()
    egitim_ogretim_yili1 = egitim_ogretim_yili.get()
    bugunun_tarihi1 = bugunun_tarihi.get()
    adi_soyadi1 = adi_soyadi.get()
    adres1 = adres.get()
    adres_ilce1 = adres_ilce.get()
    adres_il1 = adres_il.get()
    belge_no1 = belge_no.get()
    tc_no1 = tc_no.get()
    dogum_yeri1 = dogum_yeri.get()
    baba_adi1 = baba_adi.get()
    kizi_oglu1 = kizi_oglu.get()
    diploma_derece1 = diploma_derece.get()
    diploma_tarihi1 = diploma_tarihi.get()
    diploma_sayisi1 = diploma_sayisi.get()
    duzenleyen_adi_soyadi1 = duzenleyen_adi_soyadi.get()
    duzenleyen_unvan1 = duzenleyen_unvan.get()
    onaylayan_adi_soyadi1 = onaylayan_adi_soyadi.get()
    onaylayan_unvan1 = onaylayan_unvan.get()

    okul_adi.delete(0,END)
    okul_ilce_il.delete(0,END)
    egitim_ogretim_yili.delete(0,END)
    bugunun_tarihi.delete(0,END)
    adi_soyadi.delete(0,END)
    adres.delete(0,END)
    adres_ilce.delete(0,END)
    adres_il.delete(0,END)    
    belge_no.delete(0,END)
    tc_no.delete(0,END)
    dogum_yeri.delete(0,END)
    baba_adi.delete(0,END)
    kizi_oglu.delete(0,END)
    diploma_derece.delete(0,END)
    diploma_tarihi.delete(0,END)
    diploma_sayisi.delete(0,END)
    duzenleyen_adi_soyadi.delete(0,END)
    duzenleyen_unvan.delete(0,END)
    onaylayan_adi_soyadi.delete(0,END)
    onaylayan_unvan.delete(0,END)

    if os.path.exists(adi_soyadi1+'.sq3')== False:
        vt1 = sqlite3.connect(adi_soyadi1+'.sq3')
        im1= vt1.cursor()
        im1.execute("""CREATE TABLE IF NOT EXISTS diploma(okuladi TEXT, okulilceil TEXT, egitimogretim TEXT, buguntarih TEXT, adisoyadi TEXT, adres TEXT, \
adresilce TEXT, adresil TEXT, belgeno TEXT, tcno TEXT, dogumyeri TEXT, babaadi TEXT, kizioglu TEXT, diplomaderece TEXT, diplomatarihi TEXT, diplomasayisi TEXT, \
duzenleyenadisoyadi TEXT, duzenleyenunvan TEXT, onaylayanadisoyadi TEXT, onaylayanunvan TEXT)""")
        im1.execute("""INSERT INTO diploma VALUES  (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",(okul_adi1, okul_ilce_il1, egitim_ogretim_yili1, bugunun_tarihi1, \
                                                                                                 adi_soyadi1, adres1, adres_ilce1, adres_il1, belge_no1, tc_no1, \
                                                                                                 dogum_yeri1, baba_adi1, kizi_oglu1, diploma_derece1, \
                                                                                                 diploma_tarihi1, diploma_sayisi1, duzenleyen_adi_soyadi1, \
                                                                                                 duzenleyen_unvan1, onaylayan_adi_soyadi1,onaylayan_unvan1, ))
        vt1.commit()

        liste.delete(0,END)

        for i in os.listdir(os.getcwd()):
            if i.endswith('.sq3'):
                liste.insert(END,i[0:-4])

    else:
        vt2 = sqlite3.connect(adi_soyadi1+'.sq3')
        im2= vt2.cursor()
        im2.execute("""CREATE TABLE IF NOT EXISTS diploma(okuladi TEXT, okulilceil TEXT, egitimogretim TEXT, buguntarih TEXT, adisoyadi TEXT, adres TEXT, \
adresilce TEXT, adresil TEXT, belgeno TEXT, tcno TEXT, dogumyeri TEXT, babaadi TEXT, kizioglu TEXT, diplomaderece TEXT, diplomatarihi TEXT, diplomasayisi TEXT, \
duzenleyenadisoyadi TEXT, duzenleyenunvan TEXT, onaylayanadisoyadi TEXT, onaylayanunvan TEXT)""")
        im2.execute("""UPDATE diploma SET  okuladi=?, okulilceil=?, egitimogretim=?, buguntarih=?, adisoyadi=?, adres=?, adresilce=?, adresil=?, belgeno=?, \
tcno=?, dogumyeri=?, babaadi=?, kizioglu=?, diplomaderece=?, diplomatarihi=?, diplomasayisi=?, duzenleyenadisoyadi=?, duzenleyenunvan=?, \
onaylayanadisoyadi=?, onaylayanunvan=?""",(okul_adi1, okul_ilce_il1, egitim_ogretim_yili1, bugunun_tarihi1, adi_soyadi1, adres1, adres_ilce1, adres_il1, \
                                           belge_no1, tc_no1, dogum_yeri1, baba_adi1, kizi_oglu1, diploma_derece1, diploma_tarihi1, diploma_sayisi1, \
                                           duzenleyen_adi_soyadi1, duzenleyen_unvan1, onaylayan_adi_soyadi1,onaylayan_unvan1, ))
        
        vt2.commit()
       
def cikti():
    liste1=liste.get(ACTIVE)

    vt = sqlite3.connect(str(liste1)+'.sq3')
    im= vt.cursor()
    im.execute(""" SELECT * FROM diploma""")
    rows = im.fetchall()
    data_str = ""
    sf = "{}{}{}{}{}{}{}{}{}{}{}{}{}{}{}{}{}{}{}{}"
    for row in rows:
        data_str += sf.format(row[0], row[1], row[2],row[3], row[4], row[5], row[6], row[7],row[8], row[9], row[10], row[11], row[12],row[13], row[14],row[15], \
                              row[16], row[17],row[18], row[19])
    
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    paragraph = document.add_paragraph(row[0]+" MÜDÜRLÜĞÜNE")
    paragraph.alignment = 1

    paragraph = document.add_paragraph("\t\t\t\t\t\t\t"+row[1]+"\n\n")

    paragraph = document.add_paragraph("\t"+row[2]+" Eğitim-Öğretim yılı sonunda almış olduğum İlkokul diplomamı zayi ettim. \
Diploma kayıt örneğimin çıkartılarak tarafıma verilmesi hususunda;")

    paragraph = document.add_paragraph("\tGereğini bilgilerinize arz ederim.")

    table = document.add_table(rows=1, cols=4)
    row1=table.add_row().cells
    row1[3].paragraphs[0].add_run(row[3]+"\n"+row[4]+"\n")
    row1[3].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER

    paragraph = document.add_paragraph("Adres: "+row[5]+"\n\t"+row[6]+"/"+row[7]+"\n")

    paragraph = document.add_paragraph("Belge No: "+row[8]+"\n\n")

    paragraph = document.add_paragraph("DİPLOMA KAYIT ÖRNEĞİ\n")
    paragraph.alignment = 1

    paragraph = document.add_paragraph("\tDilekçe sahibi "+row[9]+" T.C. Kimlik No.lu "+row[10]+" doğumlu "+row[11]+" "+row[12]+" "+row[4]+"' in "+row[0]+ \
                                       " 'ndan -"+row[13]+"- derece ile "+row[14]+" tarih ve "+row[15]+" sayılı diplomayı almaya hak kazandığı resmî kayıtların \
incelenmesinden anlaşılmıştır.\n")

    paragraph = document.add_paragraph("\t\t\t\t\t\t\tKAYITLARIMIZA UYGUNDUR")

    table = document.add_table(rows=0, cols=3)
    row2=table.add_row().cells
    row2[2].paragraphs[0].add_run(row[3]+"\n\n")
    row2[2].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER

    table = document.add_table(rows=0, cols=3)
    row3=table.add_row().cells
    row3[0].paragraphs[0].add_run(row[16]+"\n"+row[17])
    row3[2].paragraphs[0].add_run(row[18]+"\n"+row[19])
    row3[0].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    row3[2].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER

    document.save('DiplomaKayitOrnegi.docx')

    os.system("libreoffice --writer DiplomaKayitOrnegi.docx")

def sil():
    data_sil=liste.get(ACTIVE)

    os.remove(data_sil+".sq3")

    kayit_sil.delete(0,END)

    liste.delete(0,END)

    for i in os.listdir(os.getcwd()):
        if i.endswith('.sq3'):
            liste.insert(END,i[0:-4])

    okul_adi.delete(0,END)
    okul_ilce_il.delete(0,END)
    egitim_ogretim_yili.delete(0,END)
    bugunun_tarihi.delete(0,END)
    adi_soyadi.delete(0,END)
    adres.delete(0,END)
    adres_ilce.delete(0,END)
    adres_il.delete(0,END)    
    belge_no.delete(0,END)
    tc_no.delete(0,END)
    dogum_yeri.delete(0,END)
    baba_adi.delete(0,END)
    kizi_oglu.delete(0,END)
    diploma_derece.delete(0,END)
    diploma_tarihi.delete(0,END)
    diploma_sayisi.delete(0,END)
    duzenleyen_adi_soyadi.delete(0,END)
    duzenleyen_unvan.delete(0,END)
    onaylayan_adi_soyadi.delete(0,END)
    onaylayan_unvan.delete(0,END)

root = Tk()
root.title("Diploma Kayıt Örneği")
root.resizable(width=FALSE ,height=FALSE)
img=PhotoImage(file='diploma.png')
root.tk.call('wm','iconphoto',root._w,img)
mainframe = ttk.Frame(root,padding='3 3 12 12')
mainframe.grid(column=0, row=0)
mainframe.columnconfigure(0, weight=1)
mainframe.rowconfigure(0, weight =1)

okul_adi = ttk.Entry(mainframe, width =50)
okul_adi.grid(column = 2, row = 0)

okul_ilce_il = ttk.Entry(mainframe, width =50)
okul_ilce_il.grid(column = 2, row = 1)

egitim_ogretim_yili = ttk.Entry(mainframe, width =50)
egitim_ogretim_yili.grid(column = 2, row = 2)

bugunun_tarihi = ttk.Entry(mainframe, width =50)
bugunun_tarihi.grid(column = 2, row = 3)

adi_soyadi = ttk.Entry(mainframe, width =50)
adi_soyadi.grid(column = 2, row = 4)

adres = ttk.Entry(mainframe, width =50)
adres.grid(column = 2, row = 5)

adres_ilce = ttk.Entry(mainframe, width =50)
adres_ilce.grid(column = 2, row = 6)

adres_il = ttk.Entry(mainframe, width =50)
adres_il.grid(column = 2, row = 7)

belge_no = ttk.Entry(mainframe, width =50)
belge_no.grid(column = 2, row = 8)

tc_no = ttk.Entry(mainframe, width =50)
tc_no.grid(column = 2, row = 9)

dogum_yeri = ttk.Entry(mainframe, width =50)
dogum_yeri.grid(column = 2, row = 10)

baba_adi = ttk.Entry(mainframe, width =50)
baba_adi.grid(column = 2, row = 11)

kizi_oglu = ttk.Entry(mainframe, width =50)
kizi_oglu.grid(column = 2, row = 12)

diploma_derece = ttk.Entry(mainframe, width =50)
diploma_derece.grid(column = 2, row = 13)

diploma_tarihi = ttk.Entry(mainframe, width =50)
diploma_tarihi.grid(column = 2, row = 14)

diploma_sayisi = ttk.Entry(mainframe, width =50)
diploma_sayisi.grid(column = 2, row = 15)

duzenleyen_adi_soyadi = ttk.Entry(mainframe, width =50)
duzenleyen_adi_soyadi.grid(column = 2, row = 16)

duzenleyen_unvan = ttk.Entry(mainframe, width =50)
duzenleyen_unvan.grid(column = 2, row = 17)

onaylayan_adi_soyadi = ttk.Entry(mainframe, width =50)
onaylayan_adi_soyadi.grid(column = 2, row = 18)

onaylayan_unvan = ttk.Entry(mainframe, width =50)
onaylayan_unvan.grid(column = 2, row = 19)


kayit_sil = ttk.Entry(mainframe, width =50)
kayit_sil.grid(column = 2, row = 21)

ttk.Label(mainframe, text ='Okul Adı').grid(column = 1, row = 0)
ttk.Label(mainframe, text ='Okulun İlçesi veya İli').grid(column = 1, row = 1)
ttk.Label(mainframe, text ='Eğitim-Öğretim Yılı').grid(column = 1, row=2)
ttk.Label(mainframe, text ='Bugünün Tarihi').grid(column = 1, row=3)
ttk.Label(mainframe, text ='Adı Soyadı').grid(column = 1, row=4)
ttk.Label(mainframe, text ='Adres').grid(column = 1, row=5)
ttk.Label(mainframe, text ='Adres İlçe').grid(column = 1, row=6)
ttk.Label(mainframe, text ='Adres İl').grid(column = 1, row=7)
ttk.Label(mainframe, text ='Belge No').grid(column = 1, row=8)
ttk.Label(mainframe, text ='T.C. Kimlik No').grid(column = 1, row=9)
ttk.Label(mainframe, text ='Doğum Yeri').grid(column = 1, row=10)
ttk.Label(mainframe, text ='Baba Adı').grid(column = 1, row=11)
ttk.Label(mainframe, text ='kızı veya oğlu').grid(column = 1, row=12)
ttk.Label(mainframe, text ='Diploma Derecesi').grid(column = 1, row=13)
ttk.Label(mainframe, text ='Diploma Tarihi').grid(column = 1, row=14)
ttk.Label(mainframe, text ='Diploma Sayısı').grid(column = 1, row=15)
ttk.Label(mainframe, text ='Düzenleyenin Adı Soyadı').grid(column = 1, row=16)
ttk.Label(mainframe, text ='Düzenleyenin Ünvanı').grid(column = 1, row=17)
ttk.Label(mainframe, text ='Onaylayanın Adı Soyadı').grid(column = 1, row=18)
ttk.Label(mainframe, text ='Onaylayanın Ünvanı').grid(column = 1, row=19)

liste = Listbox(mainframe,width=70)
liste.grid(column=3, row=0,rowspan=30,  sticky=(N,S,E,W))
liste.bind("<Double-Button-1>",bilgi_girişi)

kaydirma = ttk.Scrollbar(mainframe, orient="vertical",command=liste.yview)
kaydirma.grid(column=4, row=0, rowspan=30,sticky='ns')

liste.config(yscrollcommand=kaydirma.set)
kaydirma.config(command=liste.yview)

for i in os.listdir(os.getcwd()):
    if i.endswith('.sq3'):
        liste.insert(END,i[0:-4])

ttk.Button(mainframe, text='Kaydet/Güncelle',command= kaydet).grid(column=1, row=20)
ttk.Button(mainframe, text='Sil', command= sil).grid(column=1, row=21)
ttk.Button(mainframe, text='LibreOffice Writer Ön İzleme', command = cikti).grid(column=1, row=22)

okul_adi.focus()

root.mainloop()    
